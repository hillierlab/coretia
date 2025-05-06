import matplotlib.pyplot as plt
import matplotlib.lines as mlines
import numpy as np
from matplotlib.colors import to_rgb
from PIL import Image
import matplotlib.legend as mlegend
legend_locations = list(mlegend.Legend.codes.keys())

plot_dpi = 300

def get_colors_for_curves(n_curves, categories, plates, plot_kw, color_key = 'nd50_colors'):
    colors = plot_kw.get('figures', {}).get(color_key, [])
    linestyles = []

    if len(colors) == n_curves:  # all curves have a color defined in TOML file
        # Take styles from TOML file or use default plain color
        linestyles = plot_kw.get('figures', {}).get('barplot_styles', ['-']*len(colors))

    elif categories is None:
        if len(colors) == 0:  # No colors specified in TOML, generate evenly spaced colors
            colors = generate_even_colors(n_curves)
            linestyles = ['-'] * n_curves  # Default linestyle
        elif isinstance(colors, list):
            if len(colors ) != n_curves:
                raise RuntimeError(f"Expected {n_curves} colors, got {len(colors)} colors")
            linestyles = ['-'] * len(colors)  # Default linestyle
    elif len(categories) == 2 and len(colors) == 2:
        # Generate colors based on the presence of category strings
        category_to_color = {categories[0]: colors[0], categories[1]: colors[1]}
        colors = []
        for plate in plates:
            for category, color in category_to_color.items():
                if category in plate:  # Check if category is present in the plate string
                    colors.append(color)
                    linestyles.append('-')  # Default linestyle
                    break  # Stop checking once a match is found
    elif len(categories) == 2 and len(colors) == 0:
        # Generate alternating colors and linestyles
        base_colors = generate_even_colors(n_curves)
        for i, plate in enumerate(plates):
            color = base_colors[int(i//2)]
            linestyle = '-' if (i % 2) == 0 else '--'
            colors.append(color)
            linestyles.append(linestyle)
    else:
        raise RuntimeError(f"Cannot determine how to color and style curves.")

    return colors, linestyles


def generate_even_colors(n, colormap='tab10'):
    # Get the colormap instance, ensure perceptual distinctiveness
    cmap = plt.get_cmap(colormap)

    # If the number of requested colors is greater than colormap size, cycle through the colormap
    colors = [cmap(i % cmap.N) for i in range(n)]

    return colors


def darken_color(color, factor=0.7):
    """Darken a color by a specified factor."""
    # Convert color to RGB if it is a named color
    if isinstance(color, str):
        color_rgb = to_rgb(color)
    else:
        color_rgb = color

    # Darken the color
    darkened_color = tuple(c * factor for c in color_rgb)
    return darkened_color


def plot_all(df, out_path, d_c, plot_kw=None, category_pair=None, plates=None, nd50_reference = None):
    from coretia.datahandling import normalize
    from coretia.bootstrap import nd50_barplot_mcmc, linear_bootstrap
    from coretia.process import feasible_pairs, replace_nd50_key_xlsname

    plot_kw = plot_kw.copy() if plot_kw else {}
    gblist = ['Lum', 'no_antibody_control', 'negative_control', d_c, 'sample', 'fbs', 'plate']
    gblist = list(set(gblist))  # when d_c == 'fbs', eliminate duplicate
    if d_c == 'serum' and 'fbs' in df.columns:
        df = df.drop(columns='fbs')
    df_raw = df.groupby(['plate']).apply(lambda x: x)

    if d_c not in ['moi']:
        df1, all_single_dilution = normalize(df, d_c)
        nd50_linear_bootstrapped = linear_bootstrap(df1, d_c, category_pair, np.mean, out_path)
        out_path1 = None if out_path is None else generate_output_name(out_path, 'allcurves', plot_kw)
        plot_curves(df1, d_c, out_path1, plot_kw, category_pair, nd50_linear_bootstrapped = nd50_linear_bootstrapped[0],
                    include_ab_free=plot_kw.get('figures',{}).get('include_ab_free', False))
    else:
        all_single_dilution = None

    # Plot raw curves
    plot_kw['ylabel'] = 'Raw Light Units (RLU)'
    op2 = None if out_path is None else out_path.with_stem(out_path.stem + "allcurves_rawlum")
    plot_curves(df_raw, d_c, op2, plot_kw, category_pair)

    bpath = None if out_path is None else generate_output_name(out_path, 'barplot', plot_kw)
    if d_c == 'dilution' or d_c == 'conc_ng':
        # Adjust nd50 with reference if available
        try:
            if nd50_reference is not None:
                nd50_linear_bootstrapped[0] = correct_nd50(df, nd50_reference, nd50_linear_bootstrapped[0])
        except:
            t=1
        if plot_kw.get('figures',{}).get('nd50_labels_override', 0) == 0:
            subjects = feasible_pairs(df, d_c)
        else:
            override_labels = plot_kw.get('figures', {}).get('nd50_labels')
            if override_labels is None:
                raise RuntimeError(f"TOML file 'nd50_labels_override' set to 1 but 'nd50_labels' is not defined.")
            subjects = override_labels
        colors, linestyles = get_colors_for_curves(len(plates), category_pair, plates, plot_kw, color_key = 'nd50_colors')
        hatch = ['-' if h1 == '--' else None for h1 in linestyles]
        try:
            fix = nd50_barplot_mcmc(subjects, nd50_linear_bootstrapped[0], nd50_linear_bootstrapped[1], d_c, colors, hatch=hatch, plot_kw=plot_kw,
                        categories=category_pair, log_base=2, nd50_thr_log=plot_kw.get('nd50_thr_log', 0.5))
            # Linear-bootstrap threshold above = 0.5, Hill-MCMC = 0.3
            if bpath is not None:
                if nd50_reference is not None:
                    bpath = bpath.with_stem(bpath.stem + "ref_adj")
                fix.savefig(bpath.with_suffix('.png'), dpi=plot_kw.get('dpi', 300))
                plt.close('all')
        except Exception as e:
            print(e)

        return all_single_dilution, nd50_linear_bootstrapped[0]

    return all_single_dilution, None


def correct_nd50(df, nd50_reference, nd50_x):
    """
    Adjusts ND50 samples and recomputed mean, CI_low, CI_high

    Parameters
    ----------
    df
    nd50_reference
    nd50_x

    Returns
    -------

    """
    from coretia.process import replace_nd50_key_xlsname

    nd50_x = list(nd50_x)
    file_keys = list(nd50_reference.keys())
    nd50_samekeys, path_mapping = replace_nd50_key_xlsname(df, nd50_x[1])
    nd50_factors = {f1: nd50_reference[file_keys[0]][0] / nd50_reference[f1][0] for f1 in file_keys}
    nd50_adjusted = {k1: np.array(v1) * nd50_factors[k1] for k1, v1 in nd50_samekeys.items()}
    # put back adjusted nd50 values to be able to plot
    nd50_x[1] = {okey: nd50_adjusted[path_mapping[okey]] for okey in nd50_x[1].keys()}
    low_high = {bk: np.percentile(bv, [2.5, 97.5]) if bv is not None and len(bv) else [np.nan, np.nan] for bk, bv in
                nd50_x[1].items()}
    mean1 = {bk: np.mean(bv) if bv is not None and len(bv) else np.nan for bk, bv in nd50_x[1].items()}
    nd50_x[0] = {k1: np.r_[mean1[k1], low_high[k1]] for k1 in nd50_x[0].keys()}
    return nd50_x


def plot_curves(df, d_c, out_path, plot_kw=None, category_pair=None, avg_f=np.mean, nd50_linear_bootstrapped = None,
                include_ab_free=None):
    from coretia.bootstrap import annotate_nd50_curve, sort_df_samples

    plot_kw = plot_kw or {}
    fkw = plot_kw.get('figures', {})
    style = plot_kw.get('figures', {}).get('allstyle', 'barlets')
    fig, ax = plt.subplots(figsize=plot_kw.get('figures',{}).get('allfigsize', (6, 6)))

    sorted_samples = sort_df_samples(df, category_pair)
    colors, linestyles = get_colors_for_curves(len(sorted_samples), category_pair, sorted_samples, plot_kw, color_key='allcurvescolors')

    if d_c == 'dilution':
        smallest = df[d_c].min()  # e.g 1/256 for 'dilution' or 0.01 for conc_ng
    elif d_c == 'conc_ng': # conc_ng, no-antibody control has concentration=0
        smallest = df[d_c][df[d_c]!=0].min()
    # for Raw luminescence plots, create a placeholder to show antibody-free control value:
    if include_ab_free or ('%' not in plot_kw.get("ylabel", '') and d_c not in ['fbs', 'serum', 'moi']):
        ab_free = smallest / 2
    else:
        ab_free = np.nan

    nd50_x_c = []
    for i, sample in enumerate(sorted_samples):
        sample_data = df.loc[sample].groupby(d_c)['Lum']

        # Compute the median and median absolute deviation (MAD) for error bands
        data_avg = sample_data.mean() if avg_f==np.mean else sample_data.median()
        x = data_avg.index.to_numpy()
        data_avg = data_avg.to_numpy()

        if include_ab_free or '%' not in plot_kw.get("ylabel", ''):  # Raw curves, need a nearby point for antibody-free no-antibody control
            if d_c =='dilution':
                x[x == 1] = ab_free
            elif d_c == 'conc_ng':
                x[x == 0] = ab_free   # 0 should not be on normalized plots (value at 0 is used for normalization)
        else:
            # Remove ab_free (i.e. 100 % control) from x and y vectors
            if d_c =='dilution':
                abf_pos = np.where(x==1)
            elif d_c == 'conc_ng':
                abf_pos = np.where(x==0)
            else:
                abf_pos = []  # There is no positive reference, no need to remove value
            x = np.delete(x, abf_pos)
            data_avg = np.delete(data_avg, abf_pos)

        # Sort x-axis for consistent plotting
        order = np.argsort(x)
        x_sorted = x[order]
        y_avg_sorted = data_avg[order]

        # Plot mean line
        ax.plot(
            x_sorted, y_avg_sorted,
            label=sample if len(sorted_samples) > 1 else None,
            color=colors[i], linestyle=linestyles[i],
        )

        if style == 'points':
            # Scatter: Flatten grouped data into x and y arrays
            for x_scatter, y_scatter in sample_data:
                if x_scatter == ab_free or x_scatter == 1 or x_scatter == 0:
                    continue
                ax.scatter(
                    np.repeat(x_scatter, len(y_scatter)), y_scatter, s=5, edgecolor=None,
                    color=colors[i], alpha=0.2, label=None
            )
        elif style == 'barlets':
            # Plot error bars showing standard deviation
            for x_group, y_group in sample_data:
                if x_group == ab_free or x_group == 1 or x_group == 0:
                    continue
                y_avg = avg_f(y_group)
                y_std = np.std(y_group)
                ax.errorbar(
                    x_group, y_avg, yerr=y_std,
                    color=colors[i], ecolor=colors[i], alpha=0.5,
                    elinewidth=1, capsize=0, label=None
                )

        # Add ND50 vertical line
        # Check where the curve crosses 0.5, only when curve is already normalized
        if '%' in plot_kw.get("ylabel", ''):
            if fkw.get('nd50_y_position_on_plot', 0.05) > 0:
                y_curve = [group.values for _, group in sample_data]
                if not include_ab_free and '%' in plot_kw.get("ylabel", '') and len(abf_pos):
                    del y_curve[abf_pos[0][0]]
                repeats = np.array([d1.shape[0] for d1 in y_curve])
                n_samp = repeats.min()
                if any(repeats != y_curve[0].shape[0]):
                    print(f"Not all points on the curve have same number of repeated samplings: {repeats}. Retaining only {n_samp} points.")
                if nd50_linear_bootstrapped is not None:
                    x5, low, high = nd50_linear_bootstrapped[sample]
                    annotate_nd50_curve(x5, low, high, ax, 0.7, colors[i], fkw)

            else:
                crossings = np.where(np.diff(np.sign(y_avg_sorted - 0.5)))[0]
                if len(crossings) > 0:  # If there are any crossings
                    # Choose the first crossing point
                    crossing_index = crossings[0]

                    # Get the x and y values at the crossing
                    x1, x2 = np.log2(x_sorted[crossing_index]), np.log2(x_sorted[crossing_index + 1])
                    y1, y2 = y_avg_sorted[crossing_index], y_avg_sorted[crossing_index + 1]

                    # Interpolate to find the exact x-coordinate where y = 0.5
                    x5 = 2**(x1 + (0.5 - y1) * (x2 - x1) / (y2 - y1))
                    nd50_x_c.append((x5, colors[i], linestyles[i]))
                    annotate_nd50_curve(x5, x5, x5, ax, 0.7, colors[i], fkw)
                else:  # store 1 as placeholder for "non-neutralizing"
                    nd50_x_c.append((1, colors[i], linestyles[i]))

    # Plot labels and title with plot_kw overrides
    plt.xlabel(plot_kw.get("xlabel", d_c).replace('\\n','\n'))
    plt.ylabel(plot_kw.get("ylabel", 'Luminescence (Lum)'))
    plt.yscale(plot_kw.get('plot',{}).get('yscale', 'linear'))
    
    ax.set_ylim(0, ax.get_ylim()[1])
    
    # Change 0.0-1.0 y axis to 0-100 (for percents):
    if '%' in plot_kw.get("ylabel", ''):
        import matplotlib.ticker as mticker
        ax.yaxis.set_major_formatter(mticker.PercentFormatter(xmax=1, symbol=""))
        format_x_axis_d_c(ab_free if include_ab_free else None, d_c, df)  # omit AB free for normalized values
        ylims = plot_kw.get('figures', {}).get('allcurves_ylim', [None, None])
        plt.ylim(ylims)
    elif d_c not in ['serum', 'moi']:
        format_x_axis_d_c(ab_free, d_c, df)
    elif d_c in ['serum', 'moi']:
        if d_c == 'moi':
            tick_positions = list(np.sort(df[d_c].dropna().unique()))
            ax.set_xscale('log', base=10)
            ax.set_xticks(tick_positions)
            hide_minor_ticks()
            ylims = plot_kw.get('figures',{}).get('allcurvesraw_ylim',[None, None])
            plt.ylim(ylims)

    ncols = max(1, len(sorted_samples) // 12 + 1)
    add_legend(category_pair, colors, ncols, plot_kw)

    # Show the plot
    plt.tight_layout()
    if 'yaxis' not in plot_kw.get('figures',{'allcurvesaxis':['yaxis']}).get('allcurvesaxis',['yaxis']):
        ax.get_yaxis().set_visible(False)
        ax.spines['left'].set_visible(False)

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    if out_path is not None:
        plt.savefig(out_path.with_suffix('.png'), dpi=plot_kw.get('dpi', 300))
        plt.close('all')


def add_legend(category_pair, colors, ncols, plot_kw):
    legend_on = plot_kw.get('figures', {}).get('allcurveslegend', ['outside', 12])
    if legend_on:
        legend_str = [lp1 for lp1 in legend_locations if lp1 in legend_on[0]]
        fontsize = [lg1 for lg1 in legend_on if isinstance(lg1, int)]
        if len(fontsize) == 0:
            fontsize = 10
        else:
            fontsize = fontsize[0]
        if legend_str:  # mpl legend strings
            try:
                plt.legend(frameon=False, loc=legend_str[0], ncol=ncols, fontsize=fontsize)
            except:
                t = 1
        elif 'outside' in legend_on:
            plt.legend(loc='upper left', bbox_to_anchor=(1, 1), frameon=False, ncol=ncols, fontsize=legend_on[1])
        if 'title' in legend_on:
            title_within_axes(plot_kw.get("title", 'AAV coretia Data'), **plot_kw)
    ct = plot_kw.get('figures', {}).get('allcurveslegend_category')
    if ct:
        # Custom legend
        legends = []
        titles = []
        bboxes = []
        if category_pair is not None:
            legends.append(zip(category_pair, ['gray'] * 2, ['-', '--']))
            titles.append(ct[0])
            bboxes.append(ct[1:3])
        acc = plot_kw.get('figures', {}).get('allcurveslegend_curves')
        if acc:  # categories defined, each color has is '-' and '--' curve:
            legends.append(zip(acc[3:], colors[::2], ['-'] * len(colors[::2])))
            titles.append(acc[0])
            bboxes.append(acc[1:3])
        titles = [t1.encode('utf-8').decode('unicode_escape') for t1 in titles]
        create_custom_legend(legends, titles=titles, bbox=bboxes)


def linear_bootstrap_zoomed(x_data, y_curve, color, out_path, fkw0=None, plot_height=6, dpi=plot_dpi):
    from coretia.bootstrap import annotate_nd50_curve, bootstrap_nd50_lin_2_dilutions

    x_order = np.argsort(x_data)
    xo = x_data[x_order]
    yo = y_curve[x_order]
    if fkw0 is None:
        fkw0 = {'nd50_y_position_on_plot': 0.05,
                'mcmc_show_nd50': 1,
                'nd50_plot_marker_size': 0,
                'hill_error_style': ['hill_CI95_bars']
                }

    fig0, ax0 = plt.subplots(figsize=(5, plot_height))
    bootstrapped_x_values, nd50_index = bootstrap_nd50_lin_2_dilutions(xo, yo, debug=ax0)
    if len(bootstrapped_x_values) == 0:  # no ND50 for this curve
        plt.close(fig0)
        return
    x5 = bootstrapped_x_values.mean()
    low, high = np.percentile(bootstrapped_x_values, [2.5, 97.5])
    ax0.set_xlim(xo[nd50_index] - 0.01, xo[nd50_index+1] + 0.01)
    mm = yo[nd50_index+1].min() - 0.05, yo[nd50_index].max() + 0.05
    if mm[0] >= mm[1]:
        plt.close(fig0)
        raise(f"Assumed decreasing curve but sample left from ND50 {mm[0]} is smaller than that to the right {mm[1]}.")
    ax0.set_ylim(*mm)
    annotate_nd50_curve(x5, low, high, ax0, 0.7, color, fkw0)
    fig0.savefig(out_path, dpi=dpi)
    plt.close(fig0)


def title_within_axes(text, ax=None, **kwargs):
    """
    Compact title. Adds a title above the axes without extending canvas vertically.
    """
    if ax is None:
        ax = plt.gca()
    text1 = text.encode('utf-8').decode('unicode_escape')
    ax.text(
        0.5, kwargs.get("title_y", 0.98), text1,
        transform=ax.transAxes,  # Position relative to the axes (normalized 0 to 1)
        ha=kwargs.get("ha", "center"), va=kwargs.get("va", "bottom"),  # Center horizontally and align bottom
        fontsize=kwargs.get("fontsize", 12)  # Adjust as needed
    )


def hide_minor_ticks():
    # Hide minor ticks from log axis
    from matplotlib.ticker import NullFormatter, NullLocator
    plt.gca().xaxis.set_minor_locator(NullLocator())  # Disable minor ticks
    plt.gca().yaxis.set_minor_locator(NullLocator())
    plt.gca().xaxis.set_minor_formatter(NullFormatter())
    plt.gca().yaxis.set_minor_formatter(NullFormatter())


def generate_output_name(out_path, plot_name, plot_kw):
    """
    Constructs name for saving plots, if available using TOML file figures.[plotname]
    Parameters
    ----------
    out_path: PosixPath('/Users/hd/PycharmProjects/coretia/coretia/data/paper/Fig2_moi/output/human_samples_aav9_dmem.pdf')
    plot_name: 'allcurves'
    plot_kw: {'xlabel': 'Serum Dilution', 'ylabel': 'Transduction Efficiency (%)', 'title': 'Human samples diluted in DMEM against AAV9', 'figures': {'allcurves': 'Fig 2a'}}


    Returns
    -------
        PosixPath('/Users/hd/PycharmProjects/coretia/coretia/data/paper/Fig2_moi/output/Fig 2a.pdf')

    """
    out_path1 = out_path.with_stem(out_path.stem + f"_{plot_name}")
    if 'figures' in plot_kw and plot_name in plot_kw['figures']:
        out_path1 = out_path1.with_name(plot_kw['figures'][plot_name]).with_suffix(out_path.suffix)
    return out_path1


def format_x_axis_d_c(ab_free, d_c, df):
    """
    Either dilution with 1/d tick text or concentration in ng/well for x axis.
    Parameters
    ----------
    ab_free
    d_c
    df

    Returns
    -------

    """
    if d_c == 'dilution':
        add_fraction_ticks(plt.gca(), 'x', log=True, reciprocal=True, ab_free=ab_free)
    elif d_c == 'conc_ng':  # conc_ng, replace 0 with 'Antibody-free'
        # Get the current tick positions and labels
        tick_positions = list(np.sort(df[d_c].dropna().unique()))
        tick_labels = [f'{c1:1.2f}' for c1 in tick_positions]
        if ab_free:  # only add ab_free for raw luminescence
            tick_positions = [ab_free] + tick_positions[1:]
            tick_labels = ['Antibody-free'] + tick_labels[1:]
        else: # remove 0 if present (that's the ab_free for conc_ng)
            if 0 in tick_positions:
                tick_positions = tick_positions[1:]
                tick_labels = tick_labels[1:]

        plt.gca().set_xscale('log', base=2)
        plt.gca().set_xticks(tick_positions)
        plt.gca().set_xticklabels(tick_labels)


def add_fraction_ticks(ax, axis='x', log=False, reciprocal=True, ab_free=None, x_data=None, y_data=None):
    """Add fraction-based ticks to an axis with optional log scaling and reciprocal formatting."""
    from matplotlib.ticker import FuncFormatter
    if log:
        getattr(ax,f"set_{axis}scale")('log', base=2)

    # Define the fraction formatter
    def fraction_formatter(x, pos):
        if x == ab_free:
            return "Antibody-free"
        if reciprocal:
            fraction = 1 / x if x != 0 else 1
        else:
            fraction = x
        return f'1/{fraction:1.0f}'

    getattr(ax,f"{axis}axis").set_major_formatter(FuncFormatter(fraction_formatter))


def create_custom_legend(label_color_style, titles=(None,None), bbox=((0, -0), (0.38, -0)), markersize=4, alpha=1):
    """
    Create a custom legend with colors for samples and line styles for methods.

    Parameters:
        label_color_style (list of list of 3-tuples):
            Outer list: number of columns of the legend
            Middle nested list: number of legend items
            Inner list: triplets of (label, color, style), e.g.

        bbox (tuple): Bounding box for the legend columns, ((0, -0), (0.38, -0)) puts to lower left corner.
    """
    # Custom legend handles for line styles (methods, abstracting color)
    legends = []
    for c1 in range(len(label_color_style)):
        column_legend_triplets = label_color_style[c1]
        handles = []
        for label, color, style in column_legend_triplets:
            handles.append(
                mlines.Line2D(
                    [],
                    [],
                    color=color,
                    linestyle=style if style != 'dot' else '',
                    marker='o' if style == 'dot' else None,
                    markersize=markersize if style == 'dot' else None,  # Set marker size for 'dot'
                    label=label, alpha=alpha,
                )
            )
        legends.append(plt.legend(handles=handles, bbox_to_anchor=bbox[c1], title=titles[c1], frameon=False,
                                  bbox_transform=plt.gcf().transFigure, loc='center', alignment='left'))
        legends[-1].get_title().set_ha('left')
        plt.gca().add_artist(legends[-1])


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def read_caption_from_docx(caption_docx):
    # Open the caption document
    doc = Document(caption_docx)
    caption_text = ""
    for para in doc.paragraphs:
        caption_text += para.text + "\n"  # Add text of each paragraph
    return doc, caption_text


def figure_as_word(out_path, caption_docx=None, scale=1.0, row_paths=None, keep_png=False, para_spacing=1, ext='png'):
    cp_doc, cp = None, None
    if caption_docx is not None:
        cp_doc, cp = read_caption_from_docx(caption_docx)

    doc = Document()
    # Set the page size to A4 (210mm x 297mm)
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 width in inches
    section.page_height = Inches(11.69)  # A4 height in inches
    page_width_in_inches = 6.5

    if row_paths is not None:  # multiple images are placed below each other
        if len(scale) != len(row_paths):
            raise ValueError('Scale and row_paths must have the same length')

        # Determine scaling factor for all rows
        row_width_inch = []
        for rp1 in row_paths:
            # Get the width in pixels and DPI
            img = Image.open(rp1)
            width_pixels = img.width
            dpi = img.info.get('dpi')
            if dpi is None:
                raise ValueError(f"DPI information is missing in the image: {rp1}")
            # Calculate the width in inches
            row_width_inch.append(width_pixels / dpi[0])
            img.close()

        page_scale_factor = page_width_in_inches / max(row_width_inch)
        wi_scaled = [w1 * page_scale_factor for w1 in row_width_inch]

        for rp1, wi1, s1 in zip(row_paths, wi_scaled, scale):
            doc.add_picture(str(rp1), width=Inches(wi1) * s1)
            doc.add_paragraph("\n")  # Add some space between rows
            if not keep_png:
                os.remove(rp1)  # delete temp png files
    else:
        doc.add_picture(str(out_path.with_suffix(f'.{ext}')), width=Inches(page_width_in_inches * scale))
        if not keep_png:
            os.remove(out_path.with_suffix(f'.{ext}'))

    # Add the caption text with formatting
    if cp_doc:  # If a caption document is provided, copy formatting
        for para_in_caption in cp_doc.paragraphs:
            # Create a new paragraph in the output document for each input paragraph
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Set the spacing after the paragraph (in points)
            para.paragraph_format.space_after = para_spacing  # para_spacing is in points

            for run in para_in_caption.runs:
                text = run.text
                new_run = para.add_run(text)

                # Check and apply bold and italic formatting
                if run.bold:
                    new_run.bold = True
                if run.italic:
                    new_run.italic = True

    doc.save(out_path.with_suffix('.docx'))


def crop_and_save_image(input_path, output_path, left, top, right, bottom):
    """
    Crops an image and saves the result.

    Args:
        input_path: Path to the input PNG image.
        output_path: Path to save the cropped image.
        left: Left coordinate of the crop rectangle.
        top: Top coordinate of the crop rectangle.
        right: Right coordinate of the crop rectangle.
        bottom: Bottom coordinate of the crop rectangle.
    """
    try:
        img = Image.open(input_path)
        cropped_img = img.crop((left, top, right, bottom))
        cropped_img.save(output_path)
        print(f"Image cropped and saved to {output_path}")
    except FileNotFoundError:
        print(f"Error: Input file not found at {input_path}")
    except Exception as e:
        print(f"An error occurred: {e}")